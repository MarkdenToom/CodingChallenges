#! Python3
"""What the program does:
Readlines from guests.txt
Open word document with custom styles
Create one invitation per page in the word document with the styles suggested in the book"""

from os import chdir
from pathlib import Path
import docx

chdir(Path.home()/r'Py3Projects\15-Working with PDF and Word Documents')

list_of_guests = open('guests.txt').readlines()
doc = docx.Document('invitationsTemplates.docx')

for guest in list_of_guests:
    doc.add_paragraph('It would be a pleasure to have the company of', 'Custom1')
    doc.add_paragraph(guest.strip(), 'Custom2')
    doc.add_paragraph('at 11010 Memory Lane on the Evening of', 'Custom1')
    doc.add_paragraph('April 1st', 'Custom3')
    doc.add_paragraph('at 7 o\'clock', 'Custom1')
    doc.add_page_break()
doc.save('Invitations.docx')
